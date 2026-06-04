"""Build /Users/malharpawar/Business_Website/questions.docx containing the
complete question bank — every lesson, every question — pulled directly
from lib/db/generated-questions.ts.

Output is organized by Semester → Lesson with clear headings and numbered
questions, each followed by options A–D, the correct answer, and the
explanation.
"""

from __future__ import annotations

import json
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

ROOT = Path("/Users/malharpawar/Business_Website")
TS_PATH = ROOT / "lib/db/generated-questions.ts"
SEED_PATH = ROOT / "lib/db/seed.ts"
OUT_PATH = ROOT / "questions.docx"

SMALL_WORDS = {"and", "of", "in", "the", "for", "to", "a"}


def slugify(s: str) -> str:
    s = s.lower()
    s = re.sub(r"\s+", "-", s)
    s = re.sub(r"[^a-z0-9-]", "", s)
    s = re.sub(r"-+", "-", s)
    return s.strip("-")


def titleize_from_filename(filename: str) -> str:
    """Mirror lib/db/seed.ts's titleizeFromFilename()."""
    s = re.sub(r"\.(pptx?|key)$", "", filename, flags=re.IGNORECASE)
    s = re.sub(r"-\d+(\s*\(\d+\))?$", "", s)
    s = re.sub(r"^Lesson\s*\d*\s*[_-]?\s*", "", s, flags=re.IGNORECASE)
    s = re.sub(r"^Lesson_", "", s, flags=re.IGNORECASE)
    s = re.sub(r"[_,]", " ", s)
    s = re.sub(r"\s+", " ", s).strip()
    out = []
    for i, w in enumerate(s.split(" ")):
        if i > 0 and w.lower() in SMALL_WORDS:
            out.append(w.lower())
        else:
            out.append(w[:1].upper() + w[1:])
    return " ".join(out).strip()


def parse_seed_manifest(src: str) -> dict[str, tuple[int, int, str, int]]:
    """Walk seed.ts's sem1Files / sem2Files manifests and produce a fallback
    title/lesson map keyed by powerpoint id: id -> (semester, lesson_num, title, src_order).
    """
    out: dict[str, tuple[int, int, str, int]] = {}
    order = 0
    for sem in (1, 2):
        block_m = re.search(
            rf"const\s+sem{sem}Files\s*=\s*\[(.*?)\]", src, re.DOTALL
        )
        if not block_m:
            continue
        for fn_m in re.finditer(r'"([^"]+)"', block_m.group(1)):
            filename = fn_m.group(1)
            num_m = re.search(r"Lesson\s*(\d+)", filename, re.IGNORECASE)
            lesson_num = int(num_m.group(1)) if num_m else 1000 + order
            title = titleize_from_filename(filename) or f"Lesson {order + 1}"
            pp_id = f"s{sem}-{slugify(title) or f'lesson-{order + 1}'}"
            out[pp_id] = (sem, lesson_num, title, order)
            order += 1
    return out

# A JS/JSON double-quoted string literal: handles \" and other backslash escapes.
STRING = r'"(?:[^"\\]|\\.)*"'

QUESTION_RE = re.compile(
    r"\{\s*"
    r"question:\s*(" + STRING + r"),\s*"
    r"option_a:\s*(" + STRING + r"),\s*"
    r"option_b:\s*(" + STRING + r"),\s*"
    r"option_c:\s*(" + STRING + r"),\s*"
    r"option_d:\s*(" + STRING + r"),\s*"
    r'correct:\s*"([ABCD])",\s*'
    r"explanation:\s*(" + STRING + r"),\s*"
    r"slide_ref:\s*(\d+)\s*"
    r"\}",
    re.DOTALL,
)

# Section comment that precedes each lesson key:
#   // ── Sem 1 · Lesson 1: Foundations of Business ───────
#     "s1-introduction": [
SECTION_RE = re.compile(
    r"//\s*[─\-]+\s*Sem\s+(\d+)\s*·\s*(.+?)\s*[─\-]+[^\n]*\n"
    r"\s*\"(s\d+-[a-z0-9-]+)\"\s*:\s*\[",
    re.MULTILINE,
)


def decode(js_str: str) -> str:
    """Decode a JS double-quoted string literal via JSON."""
    return json.loads(js_str)


def parse_titles(src: str) -> dict[str, tuple[int, int, str, int]]:
    """Return {key: (semester, lesson_num, title, source_order)}."""
    titles: dict[str, tuple[int, int, str, int]] = {}
    for idx, m in enumerate(SECTION_RE.finditer(src)):
        sem = int(m.group(1))
        raw_title = m.group(2).strip()
        key = m.group(3)
        lesson_m = re.match(
            r"Lesson\s*(\d+)\s*[:.\-—]?\s*(.+)", raw_title, re.IGNORECASE
        )
        if lesson_m:
            lesson_num = int(lesson_m.group(1))
            clean = lesson_m.group(2).strip()
        else:
            # No "Lesson N" prefix — use a synthetic order that keeps section order
            lesson_num = 1000 + idx
            clean = raw_title
        titles[key] = (sem, lesson_num, clean, idx)
    return titles


def parse_bank(src: str) -> dict[str, list[dict]]:
    """Walk the source, collect questions grouped by lesson key.

    Strategy: locate each `"<key>": [` start, then scan forward for question
    objects with QUESTION_RE until we reach the *next* lesson key or end of
    file. This avoids brittle bracket matching across multi-line entries.
    """
    bank: dict[str, list[dict]] = {}

    # All lesson-key start positions, in source order.
    key_re = re.compile(r'"(s\d+-[a-z0-9-]+)"\s*:\s*\[')
    matches = list(key_re.finditer(src))

    for i, m in enumerate(matches):
        key = m.group(1)
        start = m.end()
        end = matches[i + 1].start() if i + 1 < len(matches) else len(src)
        block = src[start:end]

        questions: list[dict] = []
        for q in QUESTION_RE.finditer(block):
            questions.append(
                {
                    "question": decode(q.group(1)),
                    "option_a": decode(q.group(2)),
                    "option_b": decode(q.group(3)),
                    "option_c": decode(q.group(4)),
                    "option_d": decode(q.group(5)),
                    "correct": q.group(6),
                    "explanation": decode(q.group(7)),
                    "slide_ref": int(q.group(8)),
                }
            )
        bank[key] = questions
    return bank


def build_doc(bank: dict[str, list[dict]], titles: dict[str, tuple[int, int, str, int]]) -> int:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Cover title
    heading = doc.add_heading("BusinessBoost — Complete Question Bank", level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run(
        "Tesla STEM Pythons · 25 multiple-choice questions per lesson, with explanations"
    ).italic = True

    total_lessons = len(bank)
    total_questions = sum(len(qs) for qs in bank.values())
    summary = doc.add_paragraph()
    summary.alignment = WD_ALIGN_PARAGRAPH.CENTER
    summary.add_run(f"{total_lessons} lessons · {total_questions} questions").italic = True

    doc.add_page_break()

    # Sort lessons: semester ascending, then lesson_num ascending, then source order
    def sort_key(k: str):
        meta = titles.get(k)
        if meta is None:
            return (9, 9999, k, 9999)
        sem, lesson_num, _, src_order = meta
        return (sem, lesson_num, src_order)

    sorted_keys = sorted(bank.keys(), key=sort_key)
    written = 0
    last_sem: int | None = None

    for key in sorted_keys:
        meta = titles.get(key)
        if meta is not None:
            sem, lesson_num, clean_title, _ = meta
            if last_sem is not None and sem != last_sem:
                doc.add_page_break()
            last_sem = sem
            lesson_label = (
                f"Lesson {lesson_num}" if lesson_num < 1000 else "Lesson —"
            )
            section_heading = f"Semester {sem} · {lesson_label}: {clean_title}"
        else:
            section_heading = key

        h = doc.add_heading(section_heading, level=1)

        # Subheading: powerpoint id (useful for traceability)
        idp = doc.add_paragraph()
        idp.add_run(f"Powerpoint ID: {key}").italic = True

        for i, q in enumerate(bank[key], start=1):
            p = doc.add_paragraph()
            p.add_run(f"{i}. ").bold = True
            p.add_run(q["question"])

            for letter in ("a", "b", "c", "d"):
                opt = doc.add_paragraph()
                opt.paragraph_format.left_indent = Pt(24)
                opt.add_run(f"{letter.upper()}. ").bold = True
                opt.add_run(q[f"option_{letter}"])

            ans = doc.add_paragraph()
            ans.add_run("Correct answer: ").bold = True
            ans.add_run(q["correct"])

            exp = doc.add_paragraph()
            exp.add_run("Explanation: ").bold = True
            exp.add_run(q["explanation"])

            doc.add_paragraph()  # spacer
            written += 1

        # New lesson on its own page for clean printing
        doc.add_page_break()

    doc.save(OUT_PATH)
    return written


def main() -> None:
    src = TS_PATH.read_text(encoding="utf-8")
    titles = parse_titles(src)
    bank = parse_bank(src)

    # Fallback: any lesson without a section comment gets its title/lesson-num
    # from seed.ts's filename manifest.
    seed_src = SEED_PATH.read_text(encoding="utf-8")
    manifest = parse_seed_manifest(seed_src)
    for key in bank:
        if key not in titles and key in manifest:
            titles[key] = manifest[key]

    # Sanity: report any lessons that yielded zero questions (parser regression alert)
    empty = [k for k, v in bank.items() if not v]
    if empty:
        print(f"WARNING: {len(empty)} lesson(s) had 0 questions parsed: {empty}")

    missing_titles = [k for k in bank if k not in titles]
    if missing_titles:
        print(f"WARNING: no title resolved for: {missing_titles}")

    written = build_doc(bank, titles)
    print(f"Wrote {written} questions across {len(bank)} lessons → {OUT_PATH}")


if __name__ == "__main__":
    main()
