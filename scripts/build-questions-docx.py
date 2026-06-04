"""Build /Users/malharpawar/Business_Website/questions.docx containing the
25 hand-written MCQs for Lesson 18 (Shark Tank / Demographics / Market
Segmentation). Pulled directly from lib/db/generated-questions.ts.
"""

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

QUESTIONS = [
    {
        "q": "Market segmentation is BEST defined as:",
        "a": "Selling the same product to every possible buyer at the same price",
        "b": "Dividing a heterogeneous market into smaller, more homogeneous groups whose members share similar needs, behaviors, or characteristics",
        "c": "Pricing differently across geographic regions for tax reasons",
        "d": "A method to track quarterly sales by region",
        "correct": "B",
        "explanation": "Segmentation breaks a broad market into subgroups with similar needs, behaviors, or characteristics so the firm can tailor its marketing mix more effectively. Pricing variation alone or revenue reporting are not segmentation.",
    },
    {
        "q": "Demographic segmentation primarily groups customers by:",
        "a": "Quantifiable population characteristics such as age, gender, income, education, occupation, and family size",
        "b": "Personality traits and lifestyle choices",
        "c": "Past purchase behavior and usage frequency",
        "d": "Region, climate, and population density",
        "correct": "A",
        "explanation": "Demographics are objective, quantifiable population attributes (age, gender, income, education, occupation, family size). Personality is psychographic; usage is behavioral; region is geographic.",
    },
    {
        "q": "Psychographic segmentation divides a market by:",
        "a": "Income brackets and occupation",
        "b": "Postal code and region",
        "c": "Lifestyle, values, attitudes, interests, and personality (often via VALS or AIO frameworks)",
        "d": "Frequency of purchase only",
        "correct": "C",
        "explanation": "Psychographics use lifestyle, values, attitudes, interests, and personality. They are less observable than demographics but typically more predictive of preference and brand fit.",
    },
    {
        "q": "Geographic segmentation is most useful when:",
        "a": "Customer needs vary meaningfully by region, climate, country, or population density",
        "b": "All customers behave identically regardless of location",
        "c": "The firm sells purely digital products with zero shipping or localization needs",
        "d": "Demographics fully explain demand on their own",
        "correct": "A",
        "explanation": "Geographic segmentation is justified when location-driven factors (climate, region, urban density, country, language) materially shift customer needs — winter coats, beach apparel, right-hand-drive vehicles.",
    },
    {
        "q": "A coffee chain divides customers into 'daily heavy users' and 'occasional weekend visitors.' This is an example of:",
        "a": "Demographic segmentation",
        "b": "Geographic segmentation",
        "c": "Psychographic segmentation",
        "d": "Behavioral segmentation by usage rate",
        "correct": "D",
        "explanation": "Behavioral segmentation groups customers by what they do — usage rate, frequency, loyalty, occasion, or benefits sought — rather than who they are. Heavy vs. light usage is the canonical example.",
    },
    {
        "q": "A target market is BEST understood as:",
        "a": "The specific subset of a total market a firm chooses to serve, around which its marketing mix is built",
        "b": "Every consumer who could conceivably buy the product",
        "c": "The competitor's existing customer base",
        "d": "The team's monthly sales quota",
        "correct": "A",
        "explanation": "A target market is the chosen subset (selected from segments) on which the firm focuses its marketing mix. It is narrower than the total addressable market and is the output of the targeting step in STP.",
    },
    {
        "q": "On Shark Tank, a founder offers 10% equity in her company in exchange for a $100,000 investment. Based on the offer, the IMPLIED post-money valuation is:",
        "a": "$100,000",
        "b": "$500,000",
        "c": "$1,000,000",
        "d": "$10,000,000",
        "correct": "C",
        "explanation": "Equity percentage = investment / post-money valuation, so post-money = $100,000 / 0.10 = $1,000,000. Pre-money would be $900,000. Sharks routinely pressure-test these implied valuations against revenue and traction.",
    },
    {
        "q": "When a Shark presses an entrepreneur with 'Who, exactly, is your customer?' she is most likely probing whether the entrepreneur has:",
        "a": "A polished pitch deck",
        "b": "A clearly defined target market that has been validated with real buyers",
        "c": "A patent on the product",
        "d": "A celebrity endorsement",
        "correct": "B",
        "explanation": "Sharks invest in execution, and execution begins with knowing the target market. Without a defined and validated customer, marketing spend is unfocused and unit economics are unprovable — a strong negative signal.",
    },
    {
        "q": "A Shark Tank pitch states 'Our target market is everyone who eats food.' The most accurate critique is:",
        "a": "This is the strongest possible position because it maximizes total addressable market",
        "b": "The lack of a defined segment means the marketing mix cannot be tailored, leaving the business unfocused and capital-inefficient",
        "c": "It guarantees high gross margins",
        "d": "It eliminates the need for product differentiation",
        "correct": "B",
        "explanation": "'Everyone' is not a target market. Without segmentation, no element of the marketing mix (product, price, place, promotion) can be optimized for actual buyer preferences, and customer acquisition cost balloons.",
    },
    {
        "q": "An entrepreneur preparing for Shark Tank wants to validate her target market BEFORE pitching. The MOST reliable approach is:",
        "a": "Read industry reports about the category",
        "b": "Ask family and friends what they think of the product",
        "c": "Run targeted beta tests or pilot sales with the proposed segment and measure actual buying behavior",
        "d": "Increase her ad budget to drive impressions",
        "correct": "C",
        "explanation": "Validated buying behavior from the proposed target segment is the strongest evidence of real demand. Family-and-friends feedback is biased, industry reports are too general, and ad spend is not validation.",
    },
    {
        "q": "A subscription streaming service groups subscribers by viewing patterns rather than by age. This is an example of:",
        "a": "Demographic segmentation",
        "b": "Psychographic segmentation",
        "c": "Geographic segmentation",
        "d": "Behavioral segmentation",
        "correct": "D",
        "explanation": "Grouping by viewing patterns — what subscribers actually do — is behavioral segmentation. Streaming firms have repeatedly found behavioral signals more predictive of preference and lifetime value than demographics.",
    },
    {
        "q": "A toy company targets parents of children aged 4–7. The PRIMARY segmentation basis is:",
        "a": "Demographic (age of child and family-life stage)",
        "b": "Psychographic",
        "c": "Geographic",
        "d": "Behavioral",
        "correct": "A",
        "explanation": "Age and family-life stage are quantifiable population attributes — classic demographic dimensions. Family-life stage in particular is strongly predictive for many product categories, especially child-related goods.",
    },
    {
        "q": "A ski-jacket brand restricts paid advertising to U.S. states with at least 30 inches of average annual snowfall. This is PRIMARILY:",
        "a": "Demographic segmentation",
        "b": "Behavioral segmentation",
        "c": "Geographic segmentation",
        "d": "Psychographic segmentation",
        "correct": "C",
        "explanation": "Restricting marketing spend by climate is geographic segmentation. The firm uses location-driven need (snowfall) as the gating variable for marketing exposure.",
    },
    {
        "q": "A coffee chain offers higher-tier rewards only to customers who buy at least 15 drinks per month. This is PRIMARILY:",
        "a": "Demographic segmentation",
        "b": "Geographic segmentation",
        "c": "Behavioral segmentation by usage rate",
        "d": "Mass marketing",
        "correct": "C",
        "explanation": "Distinguishing customers by usage rate (heavy vs. light) is the textbook behavioral segmentation. Loyalty programs typically tier rewards by usage to retain heavy users — who often deliver disproportionate revenue (the 80/20 dynamic).",
    },
    {
        "q": "A two-person startup with $30,000 in capital is choosing a market strategy. The MOST defensible approach is:",
        "a": "Niche / concentrated marketing — focus on one well-defined segment to compete with limited resources",
        "b": "Mass marketing — go after the largest possible market immediately",
        "c": "Differentiated marketing — multiple full marketing mixes for multiple segments",
        "d": "Avoid all segmentation until significant revenue is reached",
        "correct": "A",
        "explanation": "With limited capital, concentrated (niche) marketing maximizes the marketing-spend-to-target-fit ratio. Mass and differentiated strategies require capital to support multiple offerings; ignoring segmentation wastes scarce dollars.",
    },
    {
        "q": "The STP marketing framework refers to:",
        "a": "Strategy, Tactics, Performance",
        "b": "Sales, Targets, Profit",
        "c": "Segmentation, Targeting, Positioning",
        "d": "Selection, Testing, Production",
        "correct": "C",
        "explanation": "STP stands for Segmentation (divide the market), Targeting (choose which segments to serve), and Positioning (decide how to be perceived in the chosen segments). It is the canonical sequence for taking a market from undefined to actionable.",
    },
    {
        "q": "The 80/20 rule (Pareto principle) as applied to a customer base typically states:",
        "a": "80% of customers produce 80% of revenue",
        "b": "Approximately 80% of revenue comes from roughly 20% of customers",
        "c": "20% of revenue comes from 20% of customers",
        "d": "Customer counts and revenue are uncorrelated",
        "correct": "B",
        "explanation": "Pareto observed that ~80% of effects come from ~20% of causes. In customer bases, a small share of heavy users typically generates the majority of revenue — a key justification for behavioral segmentation, retention investment, and loyalty programs.",
    },
    {
        "q": "Compared to demographic data, psychographic data is generally:",
        "a": "Cheaper to collect and more directly observable",
        "b": "More predictive of preference and brand fit, but harder to measure",
        "c": "Always less predictive of brand choice than demographics",
        "d": "Identical to demographic data in marketing value",
        "correct": "B",
        "explanation": "Psychographics typically predict preference better than demographics — two 35-year-olds with the same income may have very different buying patterns — but they are harder to observe directly and require surveys, panels, or behavioral inference.",
    },
    {
        "q": "A firm using MASS marketing differs from a firm using DIFFERENTIATED marketing in that:",
        "a": "Mass marketing requires more segmentation work",
        "b": "Differentiated marketing offers a single product to all customers",
        "c": "Mass marketing is restricted to luxury goods",
        "d": "Mass marketing offers a single marketing mix to the whole market, while differentiated marketing develops tailored mixes for multiple segments",
        "correct": "D",
        "explanation": "Mass marketing pursues the entire market with one offering (e.g., classic Coca-Cola). Differentiated marketing develops separate marketing mixes for multiple segments (e.g., Toyota's Camry, Lexus, and Tundra brands). Each strategy has its own cost-coverage tradeoff.",
    },
    {
        "q": "Niche marketing (also called concentrated marketing) is BEST described as:",
        "a": "Focusing all marketing effort on serving a small, well-defined segment exceptionally well",
        "b": "Selling a wide product line to every possible buyer",
        "c": "Random sampling of customers",
        "d": "Buying advertising in every available channel",
        "correct": "A",
        "explanation": "Niche marketing concentrates resources on one narrow segment, accepting smaller TAM in exchange for superior fit, lower competition, and higher margins. It is often the right strategy for early-stage firms or specialty brands.",
    },
    {
        "q": "A marketer should add a NEW segmentation dimension only when:",
        "a": "A senior executive prefers the new dimension",
        "b": "It has a memorable acronym",
        "c": "It is mentioned in industry trade press",
        "d": "It materially changes one or more decisions in the marketing mix (product, price, place, promotion)",
        "correct": "D",
        "explanation": "A useful segmentation dimension is one that changes what you do. If splitting customers along the new dimension produces the same marketing mix, the dimension adds complexity without value.",
    },
    {
        "q": "Per Kotler's classic criteria, an effective market segment should be:",
        "a": "Large, profitable, and easy to ignore",
        "b": "Measurable, substantial, accessible, differentiable, and actionable",
        "c": "Identical to a demographic group",
        "d": "Built solely from behavioral data",
        "correct": "B",
        "explanation": "Kotler's standard criteria — measurable, substantial, accessible, differentiable, actionable — are the test for whether a candidate segment is worth pursuing. A segment failing any of these is structurally weak and should be redesigned or discarded.",
    },
    {
        "q": "Treating 'millennials' or 'Gen Z' as a single homogeneous customer segment is a common mistake because:",
        "a": "Generational labels are illegal in marketing",
        "b": "Generational cohorts have identical preferences by definition",
        "c": "Demographic cohorts contain wide variation in values, behaviors, and willingness to pay; treating them as homogeneous masks the real segmentation",
        "d": "Generations are too young to have purchasing power",
        "correct": "C",
        "explanation": "Generational labels collapse millions of people into a single bucket and erase real psychographic and behavioral variation. Effective segmentation typically pairs demographic anchors with psychographic or behavioral overlays.",
    },
    {
        "q": "A frequent failure mode of OVER-segmentation is:",
        "a": "Stronger brand awareness",
        "b": "Lower customer acquisition cost",
        "c": "Larger total addressable market",
        "d": "Segments so small that the cost of designing and serving a tailored marketing mix exceeds the revenue each segment can generate",
        "correct": "D",
        "explanation": "Over-segmentation creates segments below economic minimum scale — each tailored marketing mix costs more to build and deliver than the resulting segment revenue. Segment count must respect operating economics.",
    },
    {
        "q": "The principal trade-off between a NARROW target segment and a BROAD target segment is:",
        "a": "A narrower segment offers tighter product-market fit but reduces the total addressable market",
        "b": "Narrower segments are always cheaper to serve and always more profitable",
        "c": "Broader segments are always more focused than narrow ones",
        "d": "Segment width is unrelated to addressable market",
        "correct": "A",
        "explanation": "Narrower segments enable more precise targeting and stronger fit, but at the cost of TAM. The decision is a portfolio question: pick the segment width where revenue × fit × competition is best, given the firm's resources.",
    },
]


def build_doc(path: str) -> None:
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_heading(
        "Lesson 18 — Shark Tank, Demographics & Market Segmentation",
        level=1,
    )
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle.add_run("25 multiple-choice questions with explanations")
    sub_run.italic = True

    doc.add_paragraph()

    for i, q in enumerate(QUESTIONS, start=1):
        p = doc.add_paragraph()
        p.add_run(f"{i}. ").bold = True
        p.add_run(q["q"])

        for letter in ("a", "b", "c", "d"):
            opt = doc.add_paragraph(style="List Bullet")
            opt.paragraph_format.left_indent = Pt(24)
            opt.add_run(f"{letter.upper()}. ").bold = True
            opt.add_run(q[letter])

        ans = doc.add_paragraph()
        ans.add_run("Correct answer: ").bold = True
        ans.add_run(q["correct"])

        exp = doc.add_paragraph()
        exp.add_run("Explanation: ").bold = True
        exp.add_run(q["explanation"])

        doc.add_paragraph()

    doc.save(path)
    print(f"Wrote {len(QUESTIONS)} questions to {path}")


if __name__ == "__main__":
    build_doc("/Users/malharpawar/Business_Website/questions.docx")
